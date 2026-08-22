#!/usr/bin/env python3
"""Gera o Word com as sugestões de implementação para o painel Gestão DMA (contrato 61/2023)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

COR_AZUL = RGBColor(0x1D, 0x3B, 0x8F)
COR_CINZA = RGBColor(0x55, 0x5F, 0x70)
COR_VERDE = RGBColor(0x15, 0x80, 0x3D)

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def titulo(texto):
    h = doc.add_heading(texto, level=1)
    for run in h.runs:
        run.font.color.rgb = COR_AZUL
    return h

def subtitulo(texto):
    h = doc.add_heading(texto, level=2)
    for run in h.runs:
        run.font.color.rgb = COR_AZUL
    return h

def paragrafo(texto, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.bold = bold
    return p

def item(texto):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(texto)
    return p

def check(texto):
    """Linha com caixa de seleção (☐) para o usuário marcar o que implementar."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('☐  ')
    r.font.size = Pt(13)
    r.font.color.rgb = COR_VERDE
    p.add_run(texto)
    return p

def nota(texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = COR_CINZA
    return p

# ═══════════════ CAPA ═══════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Gestão DMA — Painel dedicado ao contrato 61/2023')
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = COR_AZUL

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Sugestões de implementação para leitura e priorização')
r.font.size = Pt(13)
r.font.color.rgb = COR_CINZA

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Use as caixas (☐) para marcar o que vamos implementar a priori.')
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = COR_CINZA

doc.add_paragraph()

# ═══════════════ 1. CONTEXTO ═══════════════
titulo('1. Contexto do contrato')
paragrafo('O contrato 61/2023 (GOINFRA) é de prestação de serviços técnicos especializados de '
          'GERENCIAMENTO no âmbito da Diretoria de Manutenção (DMA), firmado com a DYNATEST '
          'Engenharia Ltda. Não é uma obra: é a contratação da própria empresa gerenciadora.')
paragrafo('O projeto (gestaoDMA) é um frontend dedicado a este único contrato, rodando em paralelo '
          'com o gemoc-frontend e usando o mesmo backend. Tudo aqui deve ser filtrado/restrito ao 61/2023.', bold=False)

subtitulo('Dados-chave (já existentes no backend)')
item(f'Valor contratado: R$ 79,15 mi | Aditivos: R$ 16,84 mi | Apostilas: R$ 9,62 mi | Total: R$ 105,6 mi')
item(f'Vigência: 13/06/2023 → 23/05/2027 (prazo prorrogável até 60 meses)')
item(f'Medido: R$ 65,46 mi | Saldo a medir: R$ 36,38 mi | Empenhado: R$ 43,53 mi | Faturado: R$ 48,46 mi | Dívida: R$ 11,78 mi')
item(f'36 medições registradas (PI + reajustes RD/RP) até a 36ª (jun/2026)')
item(f'17 gestores/fiscais registrados | Processo SEI 1706/2023')

subtitulo('Estrutura contratual (definida no contrato)')
item('10 Produtos: 8 mensais (48 meses) + 2 sob demanda (Assessoramento Especializado e Deslocamentos/Hospedagens)')
item('Hierarquia: Produto → Ação → Atividade → Subatividade')
item('Avaliação mensal por notas: Prazo, Forma e Argumento (0 / 0,3 / 0,5 / 1)')
item('Fatores de ponderação: FPD (desempenho), FPA (alocação) e FPS (senioridade)')
item('Valor Final Mensal (VFM) = Valor Ofertado × FPD × FPA × FPS')
item('Monitoramento do deságio global (VMA vs VE) com retenção cautelar')

# ═══════════════ 2. SUGESTÕES ═══════════════
titulo('2. Sugestões de implementação')

subtitulo('2.1 Visão geral do contrato (Dashboard dedicado)')
check('Cartões de visão geral: valor total, vigência, dias restantes, % medido, saldo a medir, dívida')
check('Barra/gráfico por Produto (1–10) com % medido de cada um vs. o planejado')
check('Alertas: saldo a medir alto no fim da vigência; próxima medição em atraso')

subtitulo('2.2 Acompanhamento de medições')
check('Gráfico mensal do valor medido (PI) + reajustes (RD/RP) ao longo das 36 medições')
check('Comparação VFM previsto vs. realizado por produto')
check('Detecção de buracos no calendário de medições (falta de medição mensal)')
check('Auditoria dos percentuais de reajuste (RD/RP) mês a mês')

subtitulo('2.3 Gestão por Produtos (módulo central — o que diferencia este contrato)')
check('Tela com os 10 produtos e suas ações/atividades/subatividades')
check('Lançamento de notas mensais (Prazo / Forma / Argumento) → Nota Final por ação → FPD')
check('Cálculo do VFM = Valor Ofertado × FPD × FPA × FPS')
check('Registro de não conformidades e planos de recuperação (obrigação da cláusula 5.40)')

subtitulo('2.4 Equipes (Quadros 16–23)')
check('Profissionais alocados vs. equipe referencial mínima por produto/ação')
check('Indicador FPA (ausências) e FPS (senioridade) com impacto na medição')

subtitulo('2.5 Gestão e Fiscalização (Quadro 25)')
check('Fiscais/gestores organizados por produto (hoje há 17 registros genéricos de "Gestor/Fiscal")')
check('Alertas de designação/portaria por produto')

subtitulo('2.6 Financeiro')
check('Empenho × faturamento × pagamento (tabela "pagamento" está vazia — avaliar alimentação)')
check('Acompanhamento da dívida (R$ 11,78 mi) e fluxo de pagamento das medições')
check('Monitoramento do deságio global (VMA vs VE) com retenção cautelar')

subtitulo('2.7 Prazos e documentos')
check('Próxima medição / prazos de subatividades (Quadro 4 — duração máxima em dias úteis)')
check('Documentos do contrato (contrato + aditivos + apostilas) anexados ao painel')
check('Vínculo com os processos SEI')

subtitulo('2.8 Base/registro')
check('Alimentar a tabela de pagamentos do contrato no backend (hoje vazia)')
check('Revisar a data de importação/consistência das medições (a 35ª veio antes da 1ª)')

# ═══════════════ 3. PRIORIZAÇÃO ═══════════════
titulo('3. Priorização sugerida')
paragrafo('Ordem recomendada de implementação, do maior valor gerencial para o administrador:')
item('1. Visão geral do contrato (Dashboard dedicado, filtrado ao 61/2023)')
item('2. Módulo de Gestão por Produtos com notas mensais (é o que diferencia este contrato)')
item('3. Acompanhamento de medições (gráfico mensal + alertas)')
item('4. Financeiro (empenho × faturamento × pagamento + dívida)')
item('5. Equipes e Gestão/Fiscalização (produto × fiscal/gestor)')
item('6. Prazos, documentos e processos SEI')

# ═══════════════ 4. DEPENDÊNCIAS ═══════════════
titulo('4. Dependências de dados')
paragrafo('Para implementar tudo, alguns dados precisam existir no backend:')
item('Pagamentos do contrato 61/2023 (tabela "pagamento" está vazia)')
item('Documentos anexos (tabela "documentos" está vazia para este contrato)')
item('Ordens de serviço (tabela "ordem_servico" está vazia)')
item('Notas de avaliação por produto/ação — não existem; podem ser criadas no próprio painel ou alimentadas via planilha/script')
nota('As notas de Prazo/Forma/Argumento e os fatores FPD/FPA/FPS são a parte mais rica do contrato '
     'e não estão no banco hoje — avaliar se o painel fará o lançamento ou se virão de planilhas.')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— Fim do documento —')
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = COR_CINZA

out = 'sugestoes-gestao-dma.docx'
doc.save(out)
print(f'OK: {out}')
